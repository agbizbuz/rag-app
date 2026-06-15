"""DOCX file parser using python-docx."""

from __future__ import annotations

from .base import BaseParser, Chunk


class DocxParser(BaseParser):
    supported_extensions = ("docx",)

    def parse(self, file) -> list[Chunk]:
        from docx import Document as DocxDocument

        try:
            doc = DocxDocument(file)
        except Exception:
            return []

        # Map XML elements to docx wrappers to preserve document flow order
        paragraphs_by_element = {p._element: p for p in doc.paragraphs}
        tables_by_element = {t._element: t for t in doc.tables}

        blocks: list[tuple[str, dict]] = []
        element_index = 0

        for child in doc.element.body.iterchildren():
            if child in paragraphs_by_element:
                p = paragraphs_by_element[child]
                text = self._clean(p.text)
                if not text:
                    continue

                # Formatting bullet/number lists
                style_name = p.style.name.lower() if p.style else ""
                if "bullet" in style_name:
                    text = f"- {text}"
                elif "number" in style_name:
                    text = f"1. {text}"

                blocks.append((text, {"paragraph": element_index}))
                element_index += 1

            elif child in tables_by_element:
                t = tables_by_element[child]
                # Convert table to markdown
                rows_text = []
                for row in t.rows:
                    row_cells = [self._clean(cell.text) for cell in row.cells]
                    if any(row_cells):
                        rows_text.append("| " + " | ".join(row_cells) + " |")

                if rows_text:
                    if len(rows_text) > 1:
                        num_cols = max(len(row.cells) for row in t.rows) if t.rows else 0
                        separator = "| " + " | ".join(["---"] * num_cols) + " |"
                        rows_text.insert(1, separator)

                    table_md = "\n".join(rows_text)
                    blocks.append((table_md, {"table": element_index}))
                    element_index += 1

        # Chunking / Merging Strategy
        chunks: list[Chunk] = []
        current_texts: list[str] = []
        current_metadata: dict = {}
        current_length = 0
        target_chunk_size = 1000

        source_name = getattr(file, "name", "")

        def flush():
            nonlocal current_texts, current_metadata, current_length
            if current_texts:
                merged_text = "\n\n".join(current_texts)
                metadata = {
                    "source": source_name,
                    **current_metadata
                }
                chunks.append(Chunk(text=merged_text, metadata=metadata))
                current_texts = []
                current_metadata = {}
                current_length = 0

        for text, meta in blocks:
            block_len = len(text)

            # If the block itself is larger than target chunk size, flush current and put it in its own chunk
            if block_len >= target_chunk_size:
                flush()
                chunks.append(
                    Chunk(
                        text=text,
                        metadata={
                            "source": source_name,
                            **meta
                        }
                    )
                )
                continue

            # If adding this block exceeds target chunk size, flush first
            if current_length + block_len > target_chunk_size and current_texts:
                flush()

            current_texts.append(text)
            current_length += block_len
            # Merge metadata keys (e.g. tracking start/end of paragraph or element indices)
            for k, v in meta.items():
                if k not in current_metadata:
                    current_metadata[k] = v
                else:
                    if isinstance(current_metadata[k], list):
                        current_metadata[k].append(v)
                    else:
                        current_metadata[k] = [current_metadata[k], v]

        flush()
        return chunks
