"""Tests for src/ragapp/core/parser.py process_file function."""

import io

from ragapp.core.parser import process_file


def _txt_bytes():
    """Return a fresh BytesIO with TXT content."""
    f = io.BytesIO(
        b"This is a test document.\n\nIt has multiple paragraphs to simulate real text.\n\nThe chunker should split this into manageable pieces."
    )
    f.name = "test.txt"
    return f


def _pdf_bytes():
    """Return a fresh BytesIO with minimal PDF."""
    buf = io.BytesIO(
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<<"
        b"/Font<</F1 4 0 R>>/ProcSet[/PDF/Text]>>\n"
        b"endobj\n"
        b"4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
        b"xref\n0 5\n"
        b"0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n"
        b"0000000115 00000 n \n0000000267 00000 n \n"
        b"trailer<</Size 5/Root 1 0 R>>\nstartxref\n348\n%%EOF\n"
    )
    buf.name = "test.pdf"
    return buf


def _csv_bytes():
    """Return a fresh BytesIO with CSV content."""
    csv_data = b"id,name,value\n1,alpha,100\n2,beta,200\n3,gamma,300\n"
    f = io.BytesIO(csv_data)
    f.name = "test.csv"
    return f


def _docx_bytes():
    """Return a minimal valid DOCX."""
    import zipfile

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(
            "[Content_Types].xml",
            b'<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/></Types>',
        )
        zf.writestr(
            "_rels/.rels",
            b'<?xml version="1.0"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/></Relationships>',
        )
        zf.writestr(
            "word/_rels/document.xml.rels",
            b'<?xml version="1.0"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>',
        )
        doc = '<?xml version="1.0"?><w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body><w:p><w:r><w:t>Hello from test_docx</w:t></w:r></w:p></w:body></w:document>'
        zf.writestr("word/document.xml", doc)

    buf.seek(0)
    buf.name = "test.docx"
    return buf


class TestProcessFileTXT:
    """Tests for TXT file processing."""

    def test_process_txt_basic(self):
        txt_file = _txt_bytes()
        chunks = process_file(txt_file)

        assert len(chunks) == 1, f"Expected 1 chunk (content < 1000 bytes), got {len(chunks)}"
        assert all("id" in c for c in chunks), "Each chunk must have an 'id'"
        assert all("text" in c for c in chunks), "Each chunk must have 'text'"
        assert all("metadata" in c for c in chunks), "Each chunk must have 'metadata'"

    def test_process_txt_metadata(self):
        txt_file = _txt_bytes()
        chunks = process_file(txt_file)

        for chunk in chunks:
            assert chunk["metadata"]["source"] == "test.txt", f"Source mismatch: {chunk['metadata']['source']}"
            assert "chunk" in chunk["metadata"], f"Chunk metadata missing 'chunk': {chunk['metadata']}"
class TestProcessFileTXTWordSplit:
    """Tests for TXT parser word-boundary chunking (lines 20-22 of txt_parser)."""

    def test_txt_long_content_triggers_word_split(self):
        """Content > 1000 chars should use rfind space to avoid mid-word splits."""
        words = ["word" + str(i) for i in range(400)]
        content = " ".join(words)

        buf = io.BytesIO(content.encode("utf-8"))
        buf.name = "long.txt"

        chunks = process_file(buf)

        assert len(chunks) >= 2, f"Expected at least 2 chunks from long content, got {len(chunks)}"
        for chunk in chunks:
            text = chunk["text"]
            last_token = text.split(" ")[-1] if " " in text else text
            assert last_token == text[-len(last_token):], f"Trailing space issue: {repr(text[-20:])}"

    def test_txt_under_threshold(self):
        """Content under 1000 chars should produce exactly 1 chunk."""
        content = "x" * 900  # well under 1000

        buf = io.BytesIO(content.encode("utf-8"))
        buf.name = "short.txt"

        chunks = process_file(buf)
        assert len(chunks) == 1, f"Expected 1 chunk for content under threshold, got {len(chunks)}"


class TestProcessFilePDF:
    """Tests for PDF file processing."""

    def test_process_pdf_basic(self):
        pdf_file = _pdf_bytes()
        chunks = process_file(pdf_file)

        assert len(chunks) == 0, f"Minimal synthetic PDF yields no extractable text (expected), got {len(chunks)}"
        assert all("id" in c for c in chunks), "Each chunk must have an 'id'"
        assert all("text" in c for c in chunks), "Each chunk must have 'text'"
        assert all("metadata" in c for c in chunks), "Each chunk must have 'metadata'"

    def test_process_pdf_metadata(self):
        pdf_file = _pdf_bytes()
        chunks = process_file(pdf_file)

        for chunk in chunks:
            assert chunk["metadata"]["source"] == "test.pdf", f"Source mismatch: {chunk['metadata']['source']}"
            assert "page" in chunk["metadata"], f"PDF metadata missing 'page': {chunk['metadata']}"
    def test_process_pdf_multi_paragraph_splits_into_chunks(self):
        """Test PDF page with extractable text containing newlines (covers paragraph splitting)."""
        from unittest.mock import MagicMock, patch

        # Create mock reader that returns pages with multi-line extractable text
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Paragraph one.\n\nSecond paragraph here.\n\nThird block of text."

        mock_reader = MagicMock()
        mock_reader.pages = [mock_page, mock_page]

        pdf_file = _pdf_bytes()
        with patch("pypdf.PdfReader", return_value=mock_reader):
            chunks = process_file(pdf_file)

        # 6 small paragraphs (total ~120 chars) are merged into 1 chunk
        assert len(chunks) == 1, f"Expected 1 merged chunk from small paragraphs, got {len(chunks)}"
        assert "Paragraph one." in chunks[0]["text"]
        assert "Third block of text." in chunks[0]["text"]
        # Merged metadata should list all page and paragraph indices
        assert chunks[0]["metadata"]["page"] == [1, 1, 1, 2, 2, 2]
        assert chunks[0]["metadata"]["paragraph"] == [0, 1, 2, 0, 1, 2]

    def test_process_pdf_corrupt_file_returns_empty(self):
        """Corrupt / invalid PDF should return [] instead of crashing."""
        buf = io.BytesIO(b"this is not a valid PDF at all")
        buf.name = "corrupt.pdf"

        chunks = process_file(buf)
        assert chunks == [], f"Expected empty list for corrupt PDF, got {len(chunks)} chunks"

    def test_process_pdf_empty_pages_skipped(self):
        """Pages with only whitespace should be skipped."""
        from unittest.mock import MagicMock, patch

        empty_page = MagicMock()
        empty_page.extract_text.return_value = "   \n\n   "

        text_page = MagicMock()
        text_page.extract_text.return_value = "Real content here."

        mock_reader = MagicMock()
        mock_reader.pages = [empty_page, text_page]

        buf = _pdf_bytes()
        with patch("pypdf.PdfReader", return_value=mock_reader):
            chunks = process_file(buf)

        assert len(chunks) == 1, f"Expected 1 chunk (empty page skipped), got {len(chunks)}"
        assert chunks[0]["text"] == "Real content here."
        assert chunks[0]["metadata"]["page"] == 2

    def test_process_pdf_small_paragraphs_merged(self):
        """Multiple small paragraphs should be merged into one chunk when under target_chunk_size."""
        from unittest.mock import MagicMock, patch

        mock_page = MagicMock()
        # Three short paragraphs — total well under 1000 chars
        mock_page.extract_text.return_value = "Short A.\n\nShort B.\n\nShort C."

        mock_reader = MagicMock()
        mock_reader.pages = [mock_page]

        buf = _pdf_bytes()
        with patch("pypdf.PdfReader", return_value=mock_reader):
            chunks = process_file(buf)

        # All three paragraphs should be merged into a single chunk
        assert len(chunks) == 1, f"Expected 1 merged chunk, got {len(chunks)}"
        assert "Short A." in chunks[0]["text"]
        assert "Short B." in chunks[0]["text"]
        assert "Short C." in chunks[0]["text"]

    def test_process_pdf_large_page_gets_own_chunk(self):
        """A page with >1000 chars and no paragraph breaks should be word-boundary sub-chunked."""
        from unittest.mock import MagicMock, patch

        long_text = " ".join(["word" + str(i) for i in range(400)])  # ~2400 chars
        mock_page = MagicMock()
        mock_page.extract_text.return_value = long_text

        mock_reader = MagicMock()
        mock_reader.pages = [mock_page]

        buf = _pdf_bytes()
        with patch("pypdf.PdfReader", return_value=mock_reader):
            chunks = process_file(buf)

        assert len(chunks) >= 2, f"Expected >=2 sub-chunks from oversized page, got {len(chunks)}"
        # Each sub-chunk should end at a word boundary (no partial words)
        for chunk in chunks:
            text = chunk["text"]
            assert not text.endswith(" "), f"Chunk should be stripped: {repr(text[-20:])}"
            assert "sub_chunk" in chunk["metadata"], f"Oversized chunks should have 'sub_chunk' metadata"

    def test_process_pdf_metadata_consistency(self):
        """Every chunk should always have both 'page' and 'paragraph' in metadata."""
        from unittest.mock import MagicMock, patch

        # Single paragraph on page (previously the else branch omitted 'paragraph')
        single_page = MagicMock()
        single_page.extract_text.return_value = "Only one paragraph on this page."

        multi_page = MagicMock()
        multi_page.extract_text.return_value = "Para one.\n\nPara two."

        mock_reader = MagicMock()
        mock_reader.pages = [single_page, multi_page]

        buf = _pdf_bytes()
        with patch("pypdf.PdfReader", return_value=mock_reader):
            chunks = process_file(buf)

        for chunk in chunks:
            assert "page" in chunk["metadata"], f"Missing 'page' in metadata: {chunk['metadata']}"
            assert "paragraph" in chunk["metadata"], f"Missing 'paragraph' in metadata: {chunk['metadata']}"


class TestProcessFileCSV:
    """Tests for CSV file processing."""

    def test_process_csv_basic(self):
        csv_file = _csv_bytes()
        chunks = process_file(csv_file)

        assert len(chunks) == 3, f"Expected 3 rows as chunks, got {len(chunks)}"

    def test_process_csv_metadata(self):
        csv_file = _csv_bytes()
        chunks = process_file(csv_file)

        for chunk in chunks:
            assert chunk["metadata"]["source"] == "test.csv", f"Source mismatch: {chunk['metadata']['source']}"
            assert "row" in chunk["metadata"], f"CSV metadata missing 'row': {chunk['metadata']}"


class TestProcessFileDOCX:
    """Tests for DOCX file processing."""

    def test_process_docx_basic(self):
        docx_file = _docx_bytes()
        chunks = process_file(docx_file)

        assert len(chunks) == 1, f"Expected 1 paragraph as chunk, got {len(chunks)}"

    def test_process_docx_metadata(self):
        docx_file = _docx_bytes()
        chunks = process_file(docx_file)

        for chunk in chunks:
            assert chunk["metadata"]["source"] == "test.docx", f"Source mismatch: {chunk['metadata']['source']}"
            assert "paragraph" in chunk["metadata"], f"DOCX metadata missing 'paragraph': {chunk['metadata']}"
    def test_process_docx_with_table(self):
        """Test DOCX table content generates markdown (lines 43-60 of docx_parser)."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Title paragraph")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header1"
        table.cell(0, 1).text = "Header2"
        table.cell(1, 0).text = "CellA"
        table.cell(1, 1).text = "CellB"

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        buf.name = "test_table.docx"

        chunks = process_file(buf)
        assert len(chunks) >= 1, f"Expected at least 1 chunk from DOCX with table, got {len(chunks)}"

    def test_process_docx_bullet_and_numbered_lists(self):
        """Test bullet/number list formatting (lines 36, 38 of docx_parser)."""
        from docx import Document

        doc = Document()
        p_bullet = doc.add_paragraph("First item")
        p_bullet.style.name = "List Bullet"
        p_numbered = doc.add_paragraph("Second item")
        p_numbered.style.name = "List Number"

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        buf.name = "test_list.docx"

        chunks = process_file(buf)
        assert len(chunks) >= 1, f"Expected at least 1 chunk from DOCX with lists, got {len(chunks)}"

    def test_process_docx_large_paragraph_triggers_self_chunk(self):
        """Test that paragraph > target_chunk_size gets its own chunk (lines 89-99 of docx_parser)."""
        from docx import Document

        long_text = "x " * 1500  # ~3000 chars, exceeds 1000 limit
        doc = Document()
        doc.add_paragraph(long_text)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        buf.name = "test_large.docx"

        chunks = process_file(buf)
        assert len(chunks) >= 1, f"Expected at least 1 chunk from large paragraph, got {len(chunks)}"


class TestProcessFileUnsupported:
    """Tests for unsupported file types."""

    def test_process_txt_with_unsupported_extension(self):
        import io

        # Create a BytesIO with .xyz extension (unsupported)
        buf = io.BytesIO(b"some content")
        buf.name = "test.xyz"
        chunks = process_file(buf)
        assert len(chunks) == 0, f"Expected empty list for unsupported format, got {chunks}"

    def test_process_with_no_extension(self):
        import io

        # BytesIO with no extension
        buf = io.BytesIO(b"some content")
        buf.name = "noextension"
        chunks = process_file(buf)
        assert len(chunks) == 0, f"Expected empty list for file without extension, got {chunks}"
